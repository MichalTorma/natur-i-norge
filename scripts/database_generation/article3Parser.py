
# %%
import logging
from docx import Document
import pandas as pd
import os
from sqlalchemy.orm import sessionmaker
from sqlalchemy import create_engine
import pandas as pd
import model
import logging
logger = logging.getLogger()
logger.setLevel(logging.INFO)

# logging.basicConfig(filename='./info.log', format='%(levelname)s:%(message)s', filemode='w+', level=logging.INFO)


database_file = 'db/v002.db'
try:
    os.remove(database_file)
except:
    pass

engine = create_engine(f'sqlite:///{database_file}')
model.Base.metadata.create_all(engine)
# %%
nin_nb = Document('data/NiN[2]AR3v211e3.docx')
nin_en = Document('data/GEB-2019-0582R_Suppl_e2.docx')


# %% Convenience functions

def getTables(first_row_contains, document):
    res = []
    for i, table in enumerate(document.tables):
        if len(table.row_cells(0)) < len(first_row_contains):
            continue
        this_first_row = [x.text for x in table.row_cells(0)]
        # print(this_first_row[0:len(first_row)] == first_row)
        if set(first_row_contains).issubset(set(this_first_row)):
            res.append((i,table))
    if len(res) > 1:
        logging.warning('Multiple tables')
        return res
    elif len(res) == 1:
        return res
    else:
        raise Exception('Table not found')

def getRows(contains, table):
    starts_with_set = set(contains)
    res = []
    for i, row in enumerate(table):
        if len(row) < len(contains):
            continue
        if starts_with_set.issubset(set(row)):
            res.append((i, row))
    if len(res) > 1:
        logging.warning('Multiple rows')
        return res
    elif len(res) == 1:
        return res
    else:
        return None

# %% Convenience function for creating list of lists from word table
def tableToList(table):
    return [[cell.text for cell in row.cells] for row in table.rows]

# %% Get english equivalent
def getEnglishEquivalent(id, table):
    for row in table:
        if row[0] == id:
            return row
    logging.warning(f'no english equivalent for {id}')
    return [id, 'NT', 'NT']

# %% get description table block from norwegian article 3

def getNorwegianDescriptionTable(id, document):
    try:
        tables = getTables([id], document)
    except:
        logging.warning(f'Norwegian description not found for {id}')
        return None
        
    if len(tables) > 1:
        if id == 'AS':
            return tableToList(tables[1][1])
        else:
            logging.error(f'multiple descriptions {tables}')
    
    for (i, table) in tables:
        table = tableToList(table)
        # print(table[0])
        if len(set(table[0])) == 5:
            return table
        else:
            continue
    
    logging.warning(f'Description not found: {id}')
    return None
#%% To obtain the actual string from description
def _getNorwegianDescriptionString(id, description):
    start_id = getRows([id],description)[0][0] + 2
    end_id = getRows(['bK'],description)
    if end_id == None:
        end_id = getRows(['bT'], description) 
    end_id = end_id[0][0]
    return description[start_id][0]

    # print(f'start: {start_id} end: {end_id}')

# desc = getNorwegianDescriptionTable('BK', nin_nb)
# desc_string = _getNorwegianDescriptionString('BK', desc)
# print(f'desc string: {desc_string}')

# %% To obtain elementary segments
def _getElementarySegments(description):
    bk_rows = getRows(['bK'], description)
    bt_rows = getRows(['bT'], description)
    print(f'bk_rows: {bk_rows}')
    print(f'bt_rows: {bt_rows}')
# %% get LEC table from norwegian article 3
lec_tables = getTables(['Kode', 'Navn', 'VM', 'ØSP', 'bK/ bT', 'RS'], nin_nb)
lec_tables = lec_tables[0:2]
lec_desc_no = []
for lec_table in lec_tables:
    lec_list = tableToList(lec_table[1])
    for lec_desc in lec_list:
        if lec_desc[0] != 'Kode':
            lec_desc_no.append(lec_desc)

# %% get LEC table from english article 3
lec_table_en = getTables(['Co','LEC','Description','PC','PV','Sc','ES','MG','MT'],nin_en)

lec_desc_en = tableToList(lec_table_en[0][1])[1:]
#%%

# %%
# set(map(lambda x: x[0], lec_desc_en)).symmetric_difference(set(map(lambda y: y[0] if y[0]!='' else y[1], lec_desc_no )))

# %% list thru every LEC in norwegian table and find relevant data
last_lec = None
for lec in lec_desc_no:
    offset = 0
    parent = None
    if lec[0] == '':
        offset = 1
        parent = last_lec
    else:
        last_lec = lec

    id = lec[0+offset]
    en_equivalent = getEnglishEquivalent(id, lec_desc_en)
    name_en = en_equivalent[1]
    description_en = en_equivalent[2]
    name_nb = lec[1+offset]
    if lec[1+offset] == lec[2+offset]:
        offset += 1
    pattern_of_variation_string = lec[2+offset]
    structuring_process_string = lec[3+offset]
    spatial_scale = lec[5+offset]
    knowledge_base_relations = lec[6+offset]
    knowledge_base_division = lec[7+offset]
    description_table_nb = getNorwegianDescriptionTable(id, nin_nb)
    if description_table_nb == None:
        logging.warning(f'Unable to find description table for: {id}')
    else:
        # print(description_table_nb)
        description_nb = _getNorwegianDescriptionString(id, description_table_nb)
        # elementary_segments = _getElementarySegments(description_table_nb)


    print(f'id: {id}')
    # print(f'name_nb: {name_nb}')
    # print(f'name_en: {name_en}')
    # print(f'description_en: {description_en}')
    # print(f'pattern_of_variation_string: {pattern_of_variation_string}')
    # print(f'structuring_process_string: {structuring_process_string}')
    # print(f'spatial_scale: {spatial_scale}')
    # print(f'knowledge_base_relations: {knowledge_base_relations}')
    # print(f'knowledge_base_division: {knowledge_base_division}')
    # print(f'description table nb: {description_table_nb}')
    # print(f'description  nb: {description_nb[0:100]}')
    print()

# %%
getTables(['SU'],nin_nb)
#%%
tableToList(nin_nb.tables[75])[0]