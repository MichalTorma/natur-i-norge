
# %%
import logging
from docx import Document
import pandas as pd

import logging
# %%
document = Document('data/NiN[2]AR3v211e3.docx')


# %%


def getTable(first_row):
    res = []
    for i, table in enumerate(document.tables):
        if len(table.row_cells(0)) < len(first_row):
            continue
        this_first_row = [x.text for x in table.row_cells(0)]
        # print(this_first_row[0:len(first_row)] == first_row)
        if set(first_row).issubset(set(this_first_row)):
            res.append((i,table))
    if len(res) > 1:
        logging.warning('Multiple tables')
        return res
    elif len(res) == 1:
        return res[0]
    else:
        raise Exception('Table not found')


# %%
lec_tables = getTable(['Kode', 'Navn', 'VM', 'ØSP', 'bK/ bT', 'RS'])


# %%
def tableToList(table):
    return [[cell.text for cell in row.cells] for row in table.rows]

# %%
len(lec_tables)
#%%
tableToList(document.tables[13])